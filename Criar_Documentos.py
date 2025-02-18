import os
import locale
import shutil
import zipfile
import pandas as pd
from docx import Document
from datetime import datetime

data_atual = datetime.now().strftime("%d-%m-%Y")
diretorio_tabela = f"RELATÓRIO_DOCUMENTAÇÃO {data_atual}.xlsx"
tabelas_documentacao = pd.read_excel(diretorio_tabela, sheet_name=None)

locale.setlocale(locale.LC_TIME, "pt_BR.utf8")

def get_modelo(diretorios_modelos, documento, contrato):
    diretorio_modelos = diretorios_modelos.get(contrato)
    mapa_modelos = {
    "NR6" : os.path.join(diretorio_modelos, "NRs - MODELOS", "NR6 - MODELO.docx"),
    "NR18" : os.path.join(diretorio_modelos, "NRs - MODELOS", "NR18 - MODELO.docx"),
    "NR33" : os.path.join(diretorio_modelos, "NRs - MODELOS", "NR33 - MODELO.docx"),
    "NR35" : os.path.join(diretorio_modelos, "NRs - MODELOS", "NR35 - MODELO.docx"),
    "OS" : os.path.join(diretorio_modelos, "OS - MODELO.docx")
}
    return mapa_modelos.get(documento, mapa_modelos["OS"])

def substituir_texto_OS(doc, marcador, novo_texto):
    for p in doc.paragraphs:
        if marcador in p.text:
            for run in p.runs:
                run.text = run.text.replace(marcador, novo_texto)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                if marcador in celula.text:
                    for p in celula.paragraphs:
                        for run in p.runs:
                            run.text = run.text.replace(marcador, novo_texto)

def substituir_texto_NRs(nome_modelo, substituicoes, diretorio_saida):
    temp_zip_path = diretorio_saida.replace(".docx", "_temp.zip")
    temp_folder = diretorio_saida.replace(".docx", "_temp")
    shutil.copy2(nome_modelo, diretorio_saida)
    with zipfile.ZipFile(diretorio_saida, 'r') as docx_zip:
        docx_zip.extractall(temp_folder)
    xml_path = os.path.join(temp_folder, "word", "document.xml")
    with open(xml_path, "r", encoding="utf-8") as file:
        xml_content = file.read()
    for marcador, novo_texto in substituicoes.items():
        xml_content = xml_content.replace(marcador, novo_texto)
    with open(xml_path, "w", encoding="utf-8") as file:
        file.write(xml_content)
    with zipfile.ZipFile(temp_zip_path, 'w', zipfile.ZIP_DEFLATED) as docx_zip:
        for root, _, files in os.walk(temp_folder):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, temp_folder)
                docx_zip.write(file_path, arcname)
    if os.path.exists(diretorio_saida):
        os.remove(diretorio_saida)
    os.rename(temp_zip_path, diretorio_saida)
    shutil.rmtree(temp_folder)

def gerar_documentos_pendentes(diretorios_modelos,contrato, nome_funcionario, funcao, cpf, admissao, documentos_pendentes):
    for documento in documentos_pendentes:
        modelo = get_modelo(diretorios_modelos, documento, contrato)

        if os.path.exists(modelo):
            caminho_saida = f"{documento} - {nome_funcionario}.docx"
            admissao_formatada = datetime.strptime(admissao, "%d/%m/%Y").strftime("%d de %B de %Y") if documento.startswith("NR") else admissao
            substituicoes = {
                "{{NOME}}": nome_funcionario,
                "{{FUNÇÃO}}": funcao,
                "{{CPF}}": cpf,
                "{{ADMISSÃO}}": admissao_formatada,
                "{{TREINAMENTO}}": admissao_formatada
            }
            if documento.startswith("NR"):
                substituir_texto_NRs(modelo, substituicoes, caminho_saida)
            else:
                doc = Document(modelo)
                for marcador, novo_texto in substituicoes.items():
                    substituir_texto_OS(doc, marcador, novo_texto)
                doc.save(caminho_saida)
            print(f"Documento criado para: {documento} - {nome_funcionario} ({contrato})")

def main(diretorios_modelos):
    for contrato, tabela_documentacao in tabelas_documentacao.items():
        tabela_documentacao = pd.read_excel(diretorio_tabela, header=1) 
        for _, row in tabela_documentacao.iterrows():
            nome_funcionario = row["FUNCIONÁRIO"]
            funcao = row["FUNÇÃO"]
            cpf = row["CPF"]
            admissao = row["ADMISSÃO"]
            documentos_pendentes = [doc for doc in tabela_documentacao.columns[4:] if row[doc] == "P"]
            if documentos_pendentes:
                gerar_documentos_pendentes(diretorios_modelos,contrato, nome_funcionario, funcao, cpf, admissao, documentos_pendentes)


# Ponto de alteração 
diretorios_modelos = {
    "OB186 - INHAÚMA": os.path.join(os.path.expanduser("~"), "CONSORCIO CONCREJATOEFFICO LOTE 1", "Central de Arquivos - QSMS","000 ATUAL - OBRA 186 - INHAÚMA","MATRIZ DE DOCUMENTOS\MODELOS"),
    "OB201 - SÃO GONÇALO": os.path.join(os.path.expanduser("~"), "CONSORCIO CONCREJATOEFFICO LOTE 1", "Central de Arquivos - QSMS", "000 ATUAL - OBRA 201 - SÃO GONÇALO", "MATRIZ DE DOCUMENTOS", "MODELOS")
}

main(diretorios_modelos)
