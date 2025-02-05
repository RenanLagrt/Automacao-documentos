import os
import locale
import shutil
import zipfile
import pandas as pd
from docx import Document
from datetime import datetime

# Diretórios principais
diretorio_modelos = os.path.join(os.path.expanduser("~"), "CONSORCIO CONCREJATOEFFICO LOTE 1", "Central de Arquivos - QSMS", "000 ATUAL - OBRA 186 - INHAÚMA", "MATRIZ DE DOCUMENTOS", "MODELOS")

# Caminho da planilha
data_atual = datetime.now().strftime("%d-%m-%Y")
diretorio_tabela = os.path.join(os.path.expanduser("~"), "CONSORCIO CONCREJATOEFFICO LOTE 1", "Central de Arquivos - QSMS", "000 ATUAL - OBRA 186 - INHAÚMA", "Documentação Funcionários",f"RELATÓRIO_DOCUMENTAÇÃO {data_atual}.xlsx")
tabela_documentacao = pd.read_excel(diretorio_tabela)

locale.setlocale(locale.LC_TIME, "pt_BR.utf8")

def substituir_texto_OS(doc, marcador, novo_texto):
    for p in doc.paragraphs:
        if marcador in p.text:
            for run in p.runs:
                run.text = run.text.replace(marcador, novo_texto)

    # Substituir dentro das tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                if marcador in celula.text:
                    for p in celula.paragraphs:
                        for run in p.runs:
                            run.text = run.text.replace(marcador, novo_texto)

def substituir_texto_NRs(nome_modelo, substituicoes, diretorio_saida):
    # Caminho temporário para o arquivo .docx
    temp_zip_path = diretorio_saida.replace(".docx", "_temp.zip")
    temp_folder = diretorio_saida.replace(".docx", "_temp")

    shutil.copy2(nome_modelo, diretorio_saida)

    with zipfile.ZipFile(diretorio_saida, 'r') as docx_zip:
        docx_zip.extractall(temp_folder)

    xml_path = os.path.join(temp_folder, "word", "document.xml")

    with open(xml_path, "r", encoding="utf-8") as file:
        xml_content = file.read()

    for marcador, novo_texto in substituicoes.items():
        xml_content = xml_content.replace(marcador, novo_texto)

    with open(xml_path, "w", encoding="utf-8") as file:
        file.write(xml_content)

    with zipfile.ZipFile(temp_zip_path, 'w', zipfile.ZIP_DEFLATED) as docx_zip:
        for root, _, files in os.walk(temp_folder):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, temp_folder)
                docx_zip.write(file_path, arcname)

    if os.path.exists(diretorio_saida):
        os.remove(diretorio_saida)

    os.rename(temp_zip_path, diretorio_saida)

    shutil.rmtree(temp_folder)

mapa_modelos = {
    "NR6" : os.path.join(diretorio_modelos, "NRs - MODELOS", "NR6 - MODELO.docx"),
    "NR18" : os.path.join(diretorio_modelos, "NRs - MODELOS", "NR18 - MODELO.docx"),
    "OS" : os.path.join(diretorio_modelos, "OS - MODELO.docx")
}

def gerar_documentos_pendentes(nome_funcionario, funcao, cpf, admissao, documentos_pendentes):
    for documento in documentos_pendentes:
        nome_modelo = mapa_modelos.get(documento, os.path.join(diretorio_modelos, "OS - MODELO.docx"))

        if os.path.exists(nome_modelo):
            diretorio_saida = f"{documento} - {nome_funcionario}.docx"

            # Converter a data se for uma NR
            if documento.startswith("NR"):
                admissao_formatada = datetime.strptime(admissao, "%d/%m/%Y").strftime("%d de %B de %Y")
                
            else:
                admissao_formatada = admissao

            substituicoes = {
                "{{NOME}}": nome_funcionario,
                "{{FUNÇÃO}}": funcao,
                "{{CPF}}": cpf,
                "{{ADMISSÃO}}": admissao_formatada,  
            }

            if documento.startswith("NR"):
                substituir_texto_NRs(nome_modelo, substituicoes, diretorio_saida)
            else:
                doc = Document(nome_modelo)

                for marcador, novo_texto in substituicoes.items():
                    substituir_texto_OS(doc, marcador, novo_texto)

                doc.save(diretorio_saida)

            print(f"Documento criado para: {documento} - {nome_funcionario}")

# Processar todos os funcionários
for _, row in tabela_documentacao.iterrows():
    nome_funcionario = row["FUNCIONÁRIO"]
    funcao = row["FUNÇÃO"]
    cpf = row["CPF"]
    admissao = row["ADMISSÃO"]

    documentos_pendentes = [doc for doc in tabela_documentacao.columns[4:] if row[doc] == "Pendente"]

    if documentos_pendentes:
        print(documentos_pendentes)
        gerar_documentos_pendentes(nome_funcionario, funcao, cpf, admissao, documentos_pendentes) 

print("Processamento concluído.")


